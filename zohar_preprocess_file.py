import docx
import io
import os
import re
import argparse

from regexes import REGEXES

class Options:
    "file processing options for one file"

    def __init__(self, path, mode, lang, n=None):
        self.path = path
        self.mode = mode
        self.lang = lang
        self.n = n

class Document:
    "Docx content and meta info"

    def __init__(self, paragraphs, article_id, asset):
        self.paragraphs = paragraphs
        self._id = article_id
        self._title = asset.title
        self.lang = asset.lang
        self.size = len(asset.content)
    
    def title(self):
        formatted = self._title.split('|')[0].split('-')[0].strip().replace(' ', '-')
        return f'{formatted}-{self._id}'

def parse_docx(content):
    "Converts docx bytes to a list of strings where each string represents a paragraph"

    with io.BytesIO(content) as fileobj:
        doc = docx.Document(fileobj)
        paragraphs = [p.text for p in doc.paragraphs]
        return paragraphs

def paragraphs(content):
    "Converts docx file to plain text where paragraphs are separated by a newline character"

    pars = parse_docx(content)
    return '\n'.join(p.replace('\n', ' ') for p in pars)

def regex_keep(output, regexes):
    "gets a string and a list of regexes and returns a set of indexes covering all of matches"

    keep = set()

    for regex in regexes:
        keep_regex = re.compile(regex, flags=re.MULTILINE)
        for match in keep_regex.finditer(output):
            keep.update(tuple(range(match.start(), match.end())))

    return keep

def regex_split(output, regexes):
    "gets a string and a list of regexes and returns a set of the last indexes of every match"

    split = set()

    for regex in regexes:
        split_regex = re.compile(regex, flags=re.MULTILINE)
        for match in split_regex.finditer(output):
            split.add(match.end() - 1)

    return split

def split_by_indexes(output, keep, split):
    """Splits a string by regexes, `keep` are a set of regexes of patterns where a split is impossible,
       `split` are regexes of patterns that precede every split"""

    indexes = [-1] + list(sorted(set(split) - set(keep))) + [len(output)]
    for i, j in zip(indexes, indexes[1:]):
        yield output[i+1:j+1].strip()

def split_sentences(output, lang):
    """Splits the string by senteces. The regexes defining the sentences are located in regex_{lang}.py files"""

    keep = regex_keep(output, REGEXES[lang].SENTENCES_KEEP, lang == 'he')
    split = regex_split(output, REGEXES[lang].SENTENCES_SPLIT)

    return '\n'.join(split_by_indexes(output, keep, split))

def split_characters(output, lang, n_chars):
    """Splits the string by characters. Avoids splitting in the middle of a word or
       in the middle of patterns defined by language specific ITEM regex"""

    assert lang in ('en', 'he')

    item = re.compile(REGEXES[lang].ITEM)
    space = re.compile(r'(\s+)')
    chunk = ''

    for line in output.split('\n'):
        if item.match(line):
            if chunk:
                yield chunk.strip()
            chunk = ''

        for word in space.split(line):
            chunk += word
            if len(chunk) > n_chars:
                yield chunk.strip()
                chunk = ''

    if chunk:
        yield chunk.strip()

def replace(output, lang):
    "Runs replacement regexes on a string"

    pairs = REGEXES[lang].REPLACE

    for pattern, repl in pairs:
        output = re.sub(pattern, repl, output, flags=re.MULTILINE)

    return output

def process(options, postfix):
    "the main logic: splits the input into chunks, runs replacement regexes and saves the output"

    for opt in options:
        with open(opt.path, 'rb') as fin:
            content = fin.read()

        output = paragraphs(content)
        if opt.mode == 'sentences':
            output = split_sentences(output, opt.lang)
        elif opt.mode == 'chars':
            output = '\n'.join(split_characters(output, opt.lang, opt.n))
        elif opt.mode == 'joined':
            output = ' '.join(output.split())

        output = replace(output, opt.lang)

        with open(opt.path + postfix, 'w') as fout:
            fout.write(output)
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("english_file")
    parser.add_argument("hebrew_file")
    parser.add_argument("--chunk", choices=['paragraphs', 'sentences', 'chars', 'joined'], default='paragraphs')
    parser.add_argument("--n_chars_en", help='number of chars in an english phrase', default=255)
    parser.add_argument("--n_chars_he", help='number of chars in a hebrew phrase', default=225)
    parser.add_argument("--postfix", help='postifx (attached to filename at the end)', default='.txt')

    args = parser.parse_args()

    opts = [Options(args.english_file, args.chunk, 'en', args.n_chars_en),
            Options(args.hebrew_file, args.chunk, 'he', args.n_chars_he)]

    process(opts, args.postfix)

if __name__ == "__main__":
    main()
